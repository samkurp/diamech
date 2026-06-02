import os
from dotenv import load_dotenv
from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
from openpyxl import load_workbook
from datetime import datetime
import uuid
import traceback
import json
import base64
from werkzeug.utils import secure_filename
from PIL import Image
import zipfile
import tempfile
from docx import Document
import io
import olefile
import re

# Загружаем переменные окружения
load_dotenv()

app = Flask(__name__, static_folder='static')
CORS(app)


# Конфигурация
class Config:
    # Supabase конфигурация
    SUPABASE_URL = os.environ.get('SUPABASE_URL', '')
    SUPABASE_KEY = os.environ.get('SUPABASE_KEY', '')

    # Шаблон Excel
    TEMPLATE_PATH = "1.xlsx"

    # Разрешенные форматы изображений
    ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp'}

    # Разрешенные форматы для схем
    ALLOWED_SCHEMA_EXTENSIONS = {'pdf'}

    # Максимальный размер файла (16MB)
    MAX_CONTENT_LENGTH = 16 * 1024 * 1024

    # Настройки сжатия изображений
    IMAGE_MAX_SIZE = 1024  # Максимальный размер стороны в пикселях
    IMAGE_QUALITY = 85  # Качество сжатия в процентах
    IMAGE_FORMAT = 'JPEG'  # Формат для сохранения


# ========== ИНИЦИАЛИЗАЦИЯ SUPABASE ==========
supabase = None
if Config.SUPABASE_URL and Config.SUPABASE_KEY:
    try:
        print("\n" + "=" * 50)
        print("🔄 ИНИЦИАЛИЗАЦИЯ SUPABASE")
        print("=" * 50)
        print(f"📌 URL: {Config.SUPABASE_URL[:50]}...")
        print(f"📌 Key length: {len(Config.SUPABASE_KEY)} символов")

        from supabase import create_client

        supabase = create_client(Config.SUPABASE_URL, Config.SUPABASE_KEY)

        # Проверка подключения
        test_query = supabase.table('drafts').select('*').limit(1).execute()
        print("✅ Supabase: успешно подключен!")
        print("✅ Таблица 'drafts' доступна")
        print("=" * 50 + "\n")

    except ImportError as e:
        print(f"❌ Ошибка импорта supabase: {e}")
        print("   Установите: pip install supabase==2.12.0")
        supabase = None
    except Exception as e:
        print(f"❌ Ошибка подключения к Supabase: {e}")
        print(f"   Тип ошибки: {type(e).__name__}")
        supabase = None
else:
    print("\n⚠️ Supabase не настроен - переменные окружения отсутствуют")
    print("   Установите SUPABASE_URL и SUPABASE_KEY в переменных окружения\n")


# ============================================

def compress_image(image_data, max_size=1024, quality=70):
    """
    Сжимает изображение до заданных размеров и качества
    """
    try:
        # Открываем изображение
        img = Image.open(io.BytesIO(image_data))

        # Конвертируем в RGB если нужно
        if img.mode in ('RGBA', 'LA', 'P'):
            img = img.convert('RGB')

        # Изменяем размер с сохранением пропорций
        img.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)

        # Сохраняем с сжатием
        output = io.BytesIO()
        img.save(output, format=Config.IMAGE_FORMAT, quality=quality, optimize=True)
        compressed_data = output.getvalue()

        print(f"🖼️ Изображение сжато: {len(image_data)} -> {len(compressed_data)} байт")
        return compressed_data

    except Exception as e:
        print(f"❌ Ошибка сжатия изображения: {e}")
        return image_data  # Возвращаем оригинал в случае ошибки


def extract_text_from_docx(file_data):
    """Извлекает текст из DOCX файла"""
    try:
        doc = Document(io.BytesIO(file_data))
        text_parts = []

        # Извлекаем текст из параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))

        # Извлекаем текст из заголовков
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(f"[Заголовок] {paragraph.text.strip()}")

            footer = section.footer
            for paragraph in footer.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(f"[Подвал] {paragraph.text.strip()}")

        return '\n\n'.join(text_parts) if text_parts else None

    except Exception as e:
        print(f"❌ Ошибка извлечения текста из DOCX: {e}")
        return None


def extract_text_from_doc(file_data):
    """Извлекает текст из DOC файла (старый формат)"""
    try:
        ole = olefile.OleFileIO(io.BytesIO(file_data))

        if not ole.exists('WordDocument'):
            return None

        # Пытаемся извлечь текст из разных потоков
        text_parts = []

        # Основной текст
        if ole.exists('WordDocument'):
            stream = ole.openstream('WordDocument')
            data = stream.read()

            # Пытаемся декодировать текст
            try:
                # Пробуем разные кодировки
                for encoding in ['utf-16', 'utf-8', 'latin-1', 'cp1251']:
                    try:
                        text = data.decode(encoding, errors='ignore')
                        # Очищаем от нечитаемых символов
                        text = re.sub(r'[^\x20-\x7Eа-яА-ЯёЁ\s\n\r]', ' ', text)
                        # Разбиваем на строки
                        lines = [line.strip() for line in text.split('\n') if line.strip()]
                        text_parts.extend(lines)
                        break
                    except:
                        continue
            except:
                pass

        # Пытаемся получить текст из потока 1Table
        if ole.exists('1Table'):
            stream = ole.openstream('1Table')
            data = stream.read()
            try:
                text = data.decode('utf-16', errors='ignore')
                text = re.sub(r'[^\x20-\x7Eа-яА-ЯёЁ\s\n\r]', ' ', text)
                lines = [line.strip() for line in text.split('\n') if line.strip()]
                text_parts.extend(lines)
            except:
                pass

        ole.close()

        return '\n\n'.join(text_parts) if text_parts else None

    except Exception as e:
        print(f"❌ Ошибка извлечения текста из DOC: {e}")
        return None


def extract_text_from_file(file_data, filename):
    """Определяет тип файла и извлекает текст"""
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

    if ext == 'docx':
        return extract_text_from_docx(file_data)
    elif ext == 'doc':
        return extract_text_from_doc(file_data)
    else:
        return None


def allowed_file(filename):
    """Проверка разрешенного формата файла"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in Config.ALLOWED_EXTENSIONS


def allowed_schema_file(filename):
    """Проверка разрешенного формата для схемы"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in Config.ALLOWED_SCHEMA_EXTENSIONS


def generate_folder_name(data):
    """Генерирует имя папки для сохранения"""
    machine_type = data.get('machineType', '').strip()
    lifting_capacity = data.get('liftingCapacity', '').strip()
    serial_number = data.get('serialNumber', 'unknown').strip().replace(' ', '_')
    return f"ML_{machine_type}{lifting_capacity}№{serial_number}"


def transliterate_machine_type(machine_type):
    """Транслитерирует тип станка для имени ZIP файла"""
    translit_map = {
        'В': 'B',
        'ВТ': 'BT',
        'ВМ': 'BM',
        'СП': 'SP',
        'ДБС': 'DBS'
    }
    return translit_map.get(machine_type, machine_type)


def generate_protocol_filename(data):
    """Генерирует имя файла протокола (кириллица)"""
    machine_type = data.get('machineType', '').strip()
    lifting_capacity = data.get('liftingCapacity', '').strip()
    serial_number = data.get('serialNumber', 'unknown').strip().replace(' ', '_')
    return f"{machine_type}{lifting_capacity}№{serial_number}.xlsx"


def generate_zip_filename(data):
    """Генерирует имя ZIP файла (латиница)"""
    machine_type = data.get('machineType', '').strip()
    lifting_capacity = data.get('liftingCapacity', '').strip()
    serial_number = data.get('serialNumber', 'unknown').strip().replace(' ', '_')

    # Транслитерируем тип станка
    machine_type_lat = transliterate_machine_type(machine_type)

    return f"{machine_type_lat}{lifting_capacity}№{serial_number}.zip"


# ========== КЛАСС ДЛЯ РАБОТЫ С SUPABASE ==========
class SupabaseDB:

    @staticmethod
    def save_history(draft_id, old_data, new_data, changed_by=None):
        """
        Сохраняет изменения в историю с умным сравнением текстовых полей
        """
        try:
            if supabase is None:
                print("⚠️ Supabase не инициализирован, история не сохраняется")
                return False

            print(f"\n{'=' * 50}")
            print(f"📝 СОХРАНЕНИЕ ИСТОРИИ ДЛЯ {draft_id}")
            print(f"{'=' * 50}")

            # Получаем отображаемое имя черновика
            draft = SupabaseDB.get_draft(draft_id)
            if not draft:
                display_name = "Неизвестный черновик"
                print("⚠️ Черновик не найден в БД")
            else:
                display_name = draft.get('display_name', 'Неизвестный черновик')
                print(f"📌 Отображаемое имя: {display_name}")

            # Сравниваем и находим изменившиеся поля
            changed_fields = {}

            # Все ключи из обоих словарей
            all_keys = set(old_data.keys()) | set(new_data.keys())
            print(f"📊 Всего полей для сравнения: {len(all_keys)}")

            for key in all_keys:
                old_value = old_data.get(key, '')
                new_value = new_data.get(key, '')

                # Пропускаем технические поля
                if key in ['draft_id', 'created_at', 'updated_at', 'id']:
                    continue

                # Приводим к строке
                old_str = str(old_value) if old_value is not None else ''
                new_str = str(new_value) if new_value is not None else ''

                # Если значения одинаковые - пропускаем
                if old_str == new_str:
                    continue

                print(f"🔍 Изменение в поле '{key}':")

                # Специальная обработка для текстовых полей (примечания)
                if key == 'notes' and len(old_str) > 0 and len(new_str) > 0:
                    # Находим дополнения к тексту
                    additions = SupabaseDB.find_text_additions(old_str, new_str)

                    if additions:
                        # Если нашли только дополнения
                        changed_fields[key] = {
                            'type': 'addition',
                            'old': old_str,
                            'new': new_str,
                            'additions': additions,
                            'full_text': new_str  # Полный текст для отображения при необходимости
                        }
                        print(f"   📝 Найдены дополнения: {additions}")
                    else:
                        # Если текст изменился кардинально
                        changed_fields[key] = {
                            'type': 'full_change',
                            'old': old_str,
                            'new': new_str
                        }
                        print(f"   🔄 Полная замена текста")
                else:
                    # Для обычных полей - просто сохраняем изменения
                    changed_fields[key] = {
                        'type': 'simple',
                        'old': old_str,
                        'new': new_str
                    }
                    print(f"   Было: '{old_str}'")
                    print(f"   Стало: '{new_str}'")

            # Если нет изменений - выходим
            if not changed_fields:
                print("ℹ️ Нет изменений для сохранения в истории")
                return True

            print(f"✅ Найдено изменений: {len(changed_fields)}")

            # Форматируем changed_fields для удобочитаемости
            formatted_fields = {}
            field_names = {
                'workType': 'Тип работы',
                'machineType': 'Тип станка',
                'liftingCapacity': 'Грузоподъемность',
                'serialNumber': 'Заводской номер',
                'customer': 'Заказчик',
                'driveType': 'Тип привода',
                'driveNumber': 'Номер привода',
                'brakeResistor': 'Тормозной резистор',
                'resistorCount': 'Кол-во резисторов',
                'electricMotor': 'Электродвигатель',
                'motorNumber': 'Номер двигателя',
                'EnginePower': 'Мощность',
                'angleSensor': 'Датчик угла',
                'angleSensorNumber': 'Номер датчика угла',
                'speedSensorNumber': 'Номер отметчика',
                'leftVibrationSensor': 'Левый датчик вибрации',
                'leftSensitivity': 'Чувствительность левого',
                'leftSensorNumber': 'Номер левого датчика',
                'rightVibrationSensor': 'Правый датчик вибрации',
                'rightSensitivity': 'Чувствительность правого',
                'rightSensorNumber': 'Номер правого датчика',
                'measuringDevice': 'Измерительный прибор',
                'measuringDeviceNumber': 'Номер прибора',
                'signalProcessor': 'Блок обработки',
                'signalProcessorNumber': 'Номер блока',
                'notes': 'Примечания',
                'machineStatus': 'Статус станка',
                'shippingDate': 'Дата отгрузки',
                'schemaUrl': 'Электрическая схема'
            }

            for key, change_info in changed_fields.items():
                field_display = field_names.get(key, key)

                if change_info['type'] == 'addition':
                    # Для дополнений показываем только добавленный текст
                    formatted_fields[field_display] = {
                        'type': 'addition',
                        'добавлено': change_info['additions'],
                        'полный_текст': change_info['full_text']  # Сохраняем для возможности показать полный
                    }
                elif change_info['type'] == 'full_change':
                    # Для полной замены показываем было/стало
                    formatted_fields[field_display] = {
                        'type': 'full_change',
                        'было': change_info['old'] if change_info['old'] else '<пусто>',
                        'стало': change_info['new'] if change_info['new'] else '<пусто>'
                    }
                else:
                    # Для простых полей
                    formatted_fields[field_display] = {
                        'type': 'simple',
                        'было': change_info['old'] if change_info['old'] else '<пусто>',
                        'стало': change_info['new'] if change_info['new'] else '<пусто>'
                    }

            # Сохраняем в историю
            history_data = {
                'draft_id': draft_id,
                'draft_display_name': display_name,
                'changed_fields': formatted_fields,
                'changed_by': changed_by or 'system',
                'created_at': datetime.now().isoformat()
            }

            print(f"💾 Сохраняем в БД: {json.dumps(history_data, indent=2, ensure_ascii=False)}")

            result = supabase.table('draft_history').insert(history_data).execute()
            print(f"✅ Результат сохранения: {result}")
            print(f"📝 История изменений сохранена для {draft_id}, изменено полей: {len(formatted_fields)}")
            print(f"{'=' * 50}\n")

            return True

        except Exception as e:
            print(f"❌ Ошибка сохранения истории: {e}")
            traceback.print_exc()
            return False

    @staticmethod
    def find_text_additions(old_text, new_text):
        """
        Находит добавленный текст в new_text по сравнению с old_text
        Возвращает список добавленных фрагментов
        """
        try:
            if not old_text or not new_text:
                return []

            additions = []

            # Если новый текст начинается со старого - значит текст дополнен в конце
            if new_text.startswith(old_text):
                addition = new_text[len(old_text):].strip()
                if addition:
                    additions.append({
                        'position': 'end',
                        'text': addition
                    })
                return additions

            # Если новый текст заканчивается старым - значит текст дополнен в начале
            if new_text.endswith(old_text):
                addition = new_text[:len(new_text) - len(old_text)].strip()
                if addition:
                    additions.append({
                        'position': 'start',
                        'text': addition
                    })
                return additions

            # Разбиваем на слова для более точного поиска
            old_words = old_text.split()
            new_words = new_text.split()

            # Находим добавленные слова
            added_words = []
            i, j = 0, 0

            while i < len(old_words) and j < len(new_words):
                if old_words[i] == new_words[j]:
                    i += 1
                    j += 1
                else:
                    # Слово добавлено
                    added_words.append(new_words[j])
                    j += 1

            # Добавляем оставшиеся слова из new_text
            while j < len(new_words):
                added_words.append(new_words[j])
                j += 1

            if added_words:
                additions.append({
                    'position': 'middle',
                    'text': ' '.join(added_words)
                })

            return additions

        except Exception as e:
            print(f"❌ Ошибка при поиске дополнений: {e}")
            return []

    @staticmethod
    def get_history(draft_id=None, page=1, per_page=50):
        """
        Получает историю изменений
        Если draft_id указан - для конкретного черновика
        Если нет - общую историю всех черновиков
        """
        try:
            if supabase is None:
                return []

            # Базовый запрос
            query = supabase.table('draft_history').select('*')

            # Фильтр по черновику если указан
            if draft_id:
                query = query.eq('draft_id', draft_id)

            # Пагинация
            offset = (page - 1) * per_page

            # Получаем данные с сортировкой по дате (сначала новые)
            response = query.order('created_at', desc=True) \
                .range(offset, offset + per_page - 1) \
                .execute()

            # Получаем общее количество для пагинации
            count_query = supabase.table('draft_history').select('count', count='exact')
            if draft_id:
                count_query = count_query.eq('draft_id', draft_id)
            count_response = count_query.execute()

            total = count_response.count if hasattr(count_response, 'count') else 0

            history = []
            for item in response.data:
                history.append({
                    'id': item.get('id'),
                    'draft_id': item.get('draft_id'),
                    'draft_display_name': item.get('draft_display_name'),
                    'changed_fields': item.get('changed_fields', {}),
                    'changed_by': item.get('changed_by'),
                    'created_at': item.get('created_at')
                })

            return {
                'items': history,
                'total': total,
                'page': page,
                'per_page': per_page,
                'total_pages': (total + per_page - 1) // per_page
            }

        except Exception as e:
            print(f"❌ Ошибка получения истории: {e}")
            return {
                'items': [],
                'total': 0,
                'page': page,
                'per_page': per_page,
                'total_pages': 0
            }

    @staticmethod
    def get_history_by_date(date, page=1, per_page=50):
        """
        Получает историю изменений за конкретную дату
        date в формате YYYY-MM-DD
        """
        try:
            if supabase is None:
                return []

            start_date = f"{date}T00:00:00"
            end_date = f"{date}T23:59:59"

            offset = (page - 1) * per_page

            response = supabase.table('draft_history') \
                .select('*') \
                .gte('created_at', start_date) \
                .lte('created_at', end_date) \
                .order('created_at', desc=True) \
                .range(offset, offset + per_page - 1) \
                .execute()

            count_response = supabase.table('draft_history') \
                .select('count', count='exact') \
                .gte('created_at', start_date) \
                .lte('created_at', end_date) \
                .execute()

            total = count_response.count if hasattr(count_response, 'count') else 0

            history = []
            for item in response.data:
                history.append({
                    'id': item.get('id'),
                    'draft_id': item.get('draft_id'),
                    'draft_display_name': item.get('draft_display_name'),
                    'changed_fields': item.get('changed_fields', {}),
                    'changed_by': item.get('changed_by'),
                    'created_at': item.get('created_at')
                })

            return {
                'items': history,
                'total': total,
                'page': page,
                'per_page': per_page,
                'total_pages': (total + per_page - 1) // per_page
            }

        except Exception as e:
            print(f"❌ Ошибка получения истории по дате: {e}")
            return {
                'items': [],
                'total': 0,
                'page': page,
                'per_page': per_page,
                'total_pages': 0
            }

    @staticmethod
    def get_all_drafts(filter_status=None):
        """Получает все черновики с фильтрацией"""
        try:
            if supabase is None:
                print("⚠️ Supabase не инициализирован")
                return []

            # Базовый запрос
            query = supabase.table('drafts').select('*')

            # Фильтрация по статусу
            if filter_status == 'shipped':
                query = query.eq('machine_status', 'Отгружен')
            elif filter_status == 'active':
                query = query.neq('machine_status', 'Отгружен')

            # Сортировка по дате обновления
            response = query.order('updated_at', desc=True).execute()

            drafts = []
            for draft in response.data:
                drafts.append({
                    'id': draft.get('id', ''),
                    'display_name': draft.get('display_name', 'Без названия'),
                    'machine_type': draft.get('data', {}).get('machineType', 'Неизвестно'),
                    'serial_number': draft.get('data', {}).get('serialNumber', 'Неизвестно'),
                    'lifting_capacity': draft.get('data', {}).get('liftingCapacity', 'Неизвестно'),
                    'work_type': draft.get('data', {}).get('workType', 'Не указан'),
                    'customer': draft.get('data', {}).get('customer', 'Не указан'),
                    'created_at': draft.get('created_at', ''),
                    'updated_at': draft.get('updated_at', ''),
                    'machine_status': draft.get('machine_status', 'Сборка'),
                    'shipping_date': draft.get('shipping_date', ''),
                    'customer_info': draft.get('customer_info', {})
                })

            print(f"📥 Загружено {len(drafts)} черновиков")
            return drafts

        except Exception as e:
            print(f"❌ Ошибка в get_all_drafts: {e}")
            return []

    @staticmethod
    def get_draft(draft_id):
        """Загружает конкретный черновик"""
        try:
            if supabase is None:
                return None

            response = supabase.table('drafts').select('*').eq('id', draft_id).execute()

            if not response.data:
                return None

            draft = response.data[0]

            # Загружаем изображения
            images_response = supabase.table('images') \
                .select('filename') \
                .eq('draft_id', draft_id) \
                .execute()

            draft['image_files'] = [img['filename'] for img in images_response.data]

            return draft

        except Exception as e:
            print(f"❌ Ошибка загрузки черновика {draft_id}: {e}")
            return None

    @staticmethod
    def save_draft(data, images=None):
        """Сохраняет новый черновик со сжатыми изображениями"""
        try:
            if supabase is None:
                return False, "Supabase не инициализирован", None

            # Генерируем ID из данных станка
            machine_type = data.get('machineType', 'Unknown')
            lifting_capacity = data.get('liftingCapacity', 'Unknown')
            serial_number = data.get('serialNumber', 'Unknown')

            draft_id = f"{machine_type}_{lifting_capacity}_{serial_number}"
            draft_id = "".join(c for c in draft_id if c.isalnum() or c in ('_', '-'))

            if not draft_id or draft_id == "_Unknown_Unknown":
                draft_id = str(uuid.uuid4())

            # Формируем отображаемое имя
            display_name = f"{machine_type}-{lifting_capacity} №{serial_number}"

            # Проверяем, существует ли уже черновик с таким ID
            existing_draft = SupabaseDB.get_draft(draft_id)

            # Подготавливаем данные
            draft_data = {
                'id': draft_id,
                'data': dict(data),
                'created_at': datetime.now().isoformat(),
                'updated_at': datetime.now().isoformat(),
                'display_name': display_name,
                'machine_status': data.get('machineStatus', 'Сборка'),
                'shipping_date': data.get('shippingDate', ''),
                'customer_info': {}
            }

            # ЕСЛИ ЧЕРНОВИК УЖЕ СУЩЕСТВУЕТ - СОХРАНЯЕМ ИСТОРИЮ
            if existing_draft:
                old_data = existing_draft.get('data', {})
                SupabaseDB.save_history(draft_id, old_data, dict(data))
                print(f"📝 История изменений сохранена при обновлении через save_draft")

            # Сохраняем в Supabase
            supabase.table('drafts').upsert(draft_data).execute()
            print(f"💾 Сохранен черновик: {draft_id}")

            # Сохраняем изображения со сжатием
            saved_images = []
            if images:
                for img in images:
                    if img and allowed_file(img.filename):
                        filename = secure_filename(img.filename)
                        img.seek(0)
                        img_data = img.read()

                        # Сжимаем изображение
                        compressed_data = compress_image(
                            img_data,
                            max_size=Config.IMAGE_MAX_SIZE,
                            quality=Config.IMAGE_QUALITY
                        )

                        # Конвертируем в base64
                        img_base64 = base64.b64encode(compressed_data).decode('utf-8')

                        # Определяем content type
                        content_type = img.content_type
                        if not content_type or content_type == 'application/octet-stream':
                            ext = filename.lower().split('.')[-1]
                            if ext in ['jpg', 'jpeg']:
                                content_type = 'image/jpeg'
                            elif ext == 'png':
                                content_type = 'image/png'
                            else:
                                content_type = f'image/{ext}'

                        try:
                            existing = supabase.table('images') \
                                .select('*') \
                                .eq('draft_id', draft_id) \
                                .eq('filename', filename) \
                                .execute()

                            if existing.data:
                                supabase.table('images') \
                                    .update({
                                    'image_data': img_base64,
                                    'content_type': content_type,
                                    'uploaded_at': datetime.now().isoformat()
                                }) \
                                    .eq('draft_id', draft_id) \
                                    .eq('filename', filename) \
                                    .execute()
                                print(f"🔄 Обновлено существующее изображение: {filename}")
                            else:
                                image_data = {
                                    'draft_id': draft_id,
                                    'filename': filename,
                                    'image_data': img_base64,
                                    'content_type': content_type,
                                    'uploaded_at': datetime.now().isoformat()
                                }
                                supabase.table('images').insert(image_data).execute()
                                print(f"🖼️ Сохранено новое изображение: {filename}")

                            saved_images.append(filename)

                        except Exception as e:
                            print(f"❌ Ошибка при сохранении изображения {filename}: {e}")

            return True, {
                'id': draft_id,
                'display_name': display_name,
                'saved_images': saved_images,
                'is_update': bool(existing_draft)
            }, None

        except Exception as e:
            print(f"❌ Ошибка сохранения: {e}")
            traceback.print_exc()
            return False, str(e), None

    @staticmethod
    def update_draft(draft_id, data, images=None):
        """Обновляет существующий черновик с сохранением истории изменений"""
        try:
            if supabase is None:
                return False, "Supabase не инициализирован"

            # Получаем текущие данные ДО обновления
            draft = SupabaseDB.get_draft(draft_id)
            if not draft:
                return False, "Черновик не найден"

            old_data = draft.get('data', {}).copy()

            # Обновляем данные
            current_data = draft.get('data', {})
            for key, value in dict(data).items():
                if value:  # Обновляем только непустые значения
                    current_data[key] = value

            # Подготавливаем обновление
            update_data = {
                'data': current_data,
                'updated_at': datetime.now().isoformat(),
                'machine_status': data.get('machineStatus', draft.get('machine_status', 'Сборка'))
            }

            # Обновляем display_name если изменились основные поля
            machine_type = data.get('machineType', current_data.get('machineType', ''))
            lifting_capacity = data.get('liftingCapacity', current_data.get('liftingCapacity', ''))
            serial_number = data.get('serialNumber', current_data.get('serialNumber', ''))

            if machine_type and lifting_capacity and serial_number:
                update_data['display_name'] = f"{machine_type}-{lifting_capacity} №{serial_number}"

            # Сохраняем
            supabase.table('drafts').update(update_data).eq('id', draft_id).execute()
            print(f"📝 Обновлен черновик: {draft_id}")

            # СОХРАНЯЕМ ИСТОРИЮ ИЗМЕНЕНИЙ
            SupabaseDB.save_history(draft_id, old_data, current_data)

            # ОБРАБАТЫВАЕМ ИЗОБРАЖЕНИЯ ПРИ ОБНОВЛЕНИИ
            saved_images = []
            if images:
                for img in images:
                    if img and allowed_file(img.filename):
                        filename = secure_filename(img.filename)
                        img.seek(0)
                        img_data = img.read()

                        # Сжимаем изображение
                        compressed_data = compress_image(
                            img_data,
                            max_size=Config.IMAGE_MAX_SIZE,
                            quality=Config.IMAGE_QUALITY
                        )

                        # Конвертируем в base64
                        img_base64 = base64.b64encode(compressed_data).decode('utf-8')

                        # Определяем content type
                        content_type = img.content_type
                        if not content_type or content_type == 'application/octet-stream':
                            ext = filename.lower().split('.')[-1]
                            if ext in ['jpg', 'jpeg']:
                                content_type = 'image/jpeg'
                            elif ext == 'png':
                                content_type = 'image/png'
                            else:
                                content_type = f'image/{ext}'

                        # Проверяем существование изображения
                        try:
                            existing = supabase.table('images') \
                                .select('*') \
                                .eq('draft_id', draft_id) \
                                .eq('filename', filename) \
                                .execute()

                            if existing.data:
                                # Обновляем существующее
                                supabase.table('images') \
                                    .update({
                                    'image_data': img_base64,
                                    'content_type': content_type,
                                    'uploaded_at': datetime.now().isoformat()
                                }) \
                                    .eq('draft_id', draft_id) \
                                    .eq('filename', filename) \
                                    .execute()
                                print(f"🔄 Обновлено изображение при обновлении: {filename}")
                            else:
                                # Вставляем новое
                                image_data = {
                                    'draft_id': draft_id,
                                    'filename': filename,
                                    'image_data': img_base64,
                                    'content_type': content_type,
                                    'uploaded_at': datetime.now().isoformat()
                                }
                                supabase.table('images').insert(image_data).execute()
                                print(f"🖼️ Добавлено новое изображение при обновлении: {filename}")

                            saved_images.append(filename)

                        except Exception as e:
                            print(f"❌ Ошибка при обновлении изображения {filename}: {e}")

            # Сохраняем файл заявки, если есть
            if 'requestFile' in data:
                request_file = data.get('requestFile')
                if request_file and hasattr(request_file, 'filename') and request_file.filename:
                    success, result = SupabaseDB.save_request_file(draft_id, request_file)
                    if not success:
                        print(f"⚠️ Ошибка сохранения файла заявки при обновлении: {result}")

            return True, {
                'draft_data': update_data,
                'saved_images': saved_images
            }

        except Exception as e:
            print(f"❌ Ошибка обновления: {e}")
            traceback.print_exc()
            return False, str(e)

    @staticmethod
    def update_customer_data(draft_id, customer_data):
        """Обновляет данные заказчика"""
        try:
            if supabase is None:
                return False, "Supabase не инициализирован"

            draft = SupabaseDB.get_draft(draft_id)
            if not draft:
                return False, "Черновик не найден"

            customer_info = draft.get('customer_info', {})
            customer_info.update(customer_data)
            customer_info['updated_at'] = datetime.now().isoformat()

            supabase.table('drafts') \
                .update({
                'customer_info': customer_info,
                'updated_at': datetime.now().isoformat()
            }) \
                .eq('id', draft_id) \
                .execute()

            print(f"👥 Обновлены данные заказчика: {draft_id}")
            return True, draft

        except Exception as e:
            print(f"❌ Ошибка обновления данных заказчика: {e}")
            return False, str(e)

    @staticmethod
    def get_customer_data(draft_id):
        """Получает данные заказчика"""
        try:
            if supabase is None:
                return None

            response = supabase.table('drafts') \
                .select('customer_info, data') \
                .eq('id', draft_id) \
                .execute()

            if not response.data:
                return None

            draft = response.data[0]
            customer_info = draft.get('customer_info', {})

            if customer_info:
                return customer_info

            # Если нет customer_info, возвращаем базовые данные
            customer_name = draft.get('data', {}).get('customer', 'Не указан')
            return {
                'customerName': customer_name,
                'originalCustomer': customer_name
            }

        except Exception as e:
            print(f"❌ Ошибка получения данных заказчика: {e}")
            return None

    @staticmethod
    def get_image(draft_id, filename):
        """Получает изображение из базы"""
        try:
            if supabase is None:
                return None

            response = supabase.table('images') \
                .select('image_data, content_type') \
                .eq('draft_id', draft_id) \
                .eq('filename', filename) \
                .execute()

            if not response.data:
                return None

            img = response.data[0]
            img_bytes = base64.b64decode(img['image_data'])
            return img_bytes, img.get('content_type', 'image/jpeg')

        except Exception as e:
            print(f"❌ Ошибка получения изображения: {e}")
            return None

    @staticmethod
    def delete_draft(draft_id):
        """Полностью удаляет черновик, его изображения и историю изменений"""
        try:
            if supabase is None:
                return False, "Supabase не инициализирован"

            print(f"\n{'=' * 50}")
            print(f"🗑️ УДАЛЕНИЕ ЧЕРНОВИКА: {draft_id}")
            print(f"{'=' * 50}")

            # 1. Удаляем все изображения черновика
            try:
                images_result = supabase.table('images') \
                    .delete() \
                    .eq('draft_id', draft_id) \
                    .execute()
                print(f"✅ Удалено изображений: {len(images_result.data) if images_result.data else 0}")
            except Exception as e:
                print(f"⚠️ Ошибка при удалении изображений: {e}")

            # 2. Удаляем файл заявки
            try:
                request_result = supabase.table('request_files') \
                    .delete() \
                    .eq('draft_id', draft_id) \
                    .execute()
                print(f"✅ Удален файл заявки: {len(request_result.data) if request_result.data else 0}")
            except Exception as e:
                print(f"⚠️ Ошибка при удалении файла заявки: {e}")

            # 3. Удаляем файл схемы
            try:
                schema_result = supabase.table('schema_files') \
                    .delete() \
                    .eq('draft_id', draft_id) \
                    .execute()
                print(f"✅ Удален файл схемы: {len(schema_result.data) if schema_result.data else 0}")
            except Exception as e:
                print(f"⚠️ Ошибка при удалении файла схемы: {e}")

            # 4. Удаляем историю изменений черновика
            try:
                history_result = supabase.table('draft_history') \
                    .delete() \
                    .eq('draft_id', draft_id) \
                    .execute()
                print(f"✅ Удалено записей истории: {len(history_result.data) if history_result.data else 0}")
            except Exception as e:
                print(f"⚠️ Ошибка при удалении истории: {e}")

            # 5. Удаляем сам черновик
            try:
                draft_result = supabase.table('drafts') \
                    .delete() \
                    .eq('id', draft_id) \
                    .execute()

                if draft_result.data:
                    print(f"✅ Черновик удален: {draft_id}")
                    print(f"{'=' * 50}\n")
                    return True, "Черновик успешно удален"
                else:
                    print(f"❌ Черновик не найден: {draft_id}")
                    print(f"{'=' * 50}\n")
                    return False, "Черновик не найден"
            except Exception as e:
                print(f"❌ Ошибка при удалении черновика: {e}")
                print(f"{'=' * 50}\n")
                return False, str(e)

        except Exception as e:
            print(f"❌ Критическая ошибка при удалении: {e}")
            traceback.print_exc()
            print(f"{'=' * 50}\n")
            return False, str(e)

    @staticmethod
    def save_request_file(draft_id, file):
        """Сохраняет файл заявки и извлекает текст"""
        try:
            if supabase is None:
                return False, "Supabase не инициализирован"

            if not file or not file.filename:
                return False, "Файл не выбран"

            filename = secure_filename(file.filename)
            ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

            allowed_extensions = {'doc', 'docx'}
            if ext not in allowed_extensions:
                return False, "Неподдерживаемый формат файла. Используйте DOC или DOCX"

            file.seek(0)
            file_data = file.read()

            # Извлекаем текст из файла
            print(f"📝 Извлечение текста из файла: {filename}")
            extracted_text = extract_text_from_file(file_data, filename)

            if extracted_text:
                print(f"✅ Текст извлечен. Длина: {len(extracted_text)} символов")
                print(f"   Первые 200 символов: {extracted_text[:200]}...")
                has_text = True
            else:
                print(f"⚠️ Не удалось извлечь текст из файла")
                extracted_text = None
                has_text = False

            # Конвертируем файл в base64
            file_base64 = base64.b64encode(file_data).decode('utf-8')

            # Определяем content type
            content_type = 'application/msword' if ext == 'doc' else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

            # Сохраняем в таблицу request_files
            request_data = {
                'draft_id': draft_id,
                'filename': filename,
                'file_data': file_base64,
                'content_type': content_type,
                'extracted_text': extracted_text,
                'has_text': has_text,
                'uploaded_at': datetime.now().isoformat()
            }

            # Проверяем, существует ли уже файл
            existing = supabase.table('request_files') \
                .select('*') \
                .eq('draft_id', draft_id) \
                .execute()

            if existing.data:
                # Обновляем существующий
                supabase.table('request_files') \
                    .update(request_data) \
                    .eq('draft_id', draft_id) \
                    .execute()
                print(f"📄 Обновлен файл заявки: {filename}")
            else:
                # Вставляем новый
                supabase.table('request_files').insert(request_data).execute()
                print(f"📄 Сохранен файл заявки: {filename}")

            return True, {
                'filename': filename,
                'has_text': has_text,
                'text_preview': extracted_text[:200] + '...' if extracted_text else None
            }

        except Exception as e:
            print(f"❌ Ошибка сохранения файла заявки: {e}")
            traceback.print_exc()
            return False, str(e)

    @staticmethod
    def get_request_file(draft_id):
        """Получает файл заявки"""
        try:
            if supabase is None:
                return None

            response = supabase.table('request_files') \
                .select('*') \
                .eq('draft_id', draft_id) \
                .execute()

            if not response.data:
                return None

            file_data = response.data[0]
            file_bytes = base64.b64decode(file_data['file_data'])
            return file_bytes, file_data['filename'], file_data['content_type']

        except Exception as e:
            print(f"❌ Ошибка получения файла заявки: {e}")
            return None

    @staticmethod
    def delete_request_file(draft_id):
        """Удаляет файл заявки"""
        try:
            if supabase is None:
                return False

            supabase.table('request_files') \
                .delete() \
                .eq('draft_id', draft_id) \
                .execute()

            print(f"📄 Удален файл заявки для {draft_id}")
            return True

        except Exception as e:
            print(f"❌ Ошибка удаления файла заявки: {e}")
            return False

    # ========== НОВЫЕ МЕТОДЫ ДЛЯ РАБОТЫ СО СХЕМОЙ ==========

    @staticmethod
    def save_schema_file(draft_id, file):
        """Сохраняет файл электрической схемы (PDF)"""
        try:
            if supabase is None:
                return False, "Supabase не инициализирован"

            if not file or not file.filename:
                return False, "Файл не выбран"

            filename = secure_filename(file.filename)
            ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

            if ext != 'pdf':
                return False, "Неподдерживаемый формат файла. Используйте PDF"

            file.seek(0)
            file_data = file.read()

            # Конвертируем файл в base64
            file_base64 = base64.b64encode(file_data).decode('utf-8')

            # Сохраняем в таблицу schema_files
            schema_data = {
                'draft_id': draft_id,
                'filename': filename,
                'file_data': file_base64,
                'content_type': 'application/pdf',
                'uploaded_at': datetime.now().isoformat()
            }

            # Проверяем, существует ли уже файл
            existing = supabase.table('schema_files') \
                .select('*') \
                .eq('draft_id', draft_id) \
                .execute()

            if existing.data:
                # Обновляем существующий
                supabase.table('schema_files') \
                    .update(schema_data) \
                    .eq('draft_id', draft_id) \
                    .execute()
                print(f"📐 Обновлен файл схемы: {filename}")
            else:
                # Вставляем новый
                supabase.table('schema_files').insert(schema_data).execute()
                print(f"📐 Сохранен файл схемы: {filename}")

            # Сохраняем URL в data черновика для совместимости
            draft = SupabaseDB.get_draft(draft_id)
            if draft:
                current_data = draft.get('data', {})
                current_data['schemaUrl'] = filename
                supabase.table('drafts') \
                    .update({'data': current_data}) \
                    .eq('id', draft_id) \
                    .execute()

            return True, filename

        except Exception as e:
            print(f"❌ Ошибка сохранения файла схемы: {e}")
            traceback.print_exc()
            return False, str(e)

    @staticmethod
    def get_schema_file(draft_id):
        """Получает файл электрической схемы"""
        try:
            if supabase is None:
                return None

            response = supabase.table('schema_files') \
                .select('*') \
                .eq('draft_id', draft_id) \
                .execute()

            if not response.data:
                return None

            file_data = response.data[0]
            file_bytes = base64.b64decode(file_data['file_data'])
            return file_bytes, file_data['filename'], file_data['content_type']

        except Exception as e:
            print(f"❌ Ошибка получения файла схемы: {e}")
            return None

    @staticmethod
    def delete_schema_file(draft_id):
        """Удаляет файл электрической схемы"""
        try:
            if supabase is None:
                return False

            supabase.table('schema_files') \
                .delete() \
                .eq('draft_id', draft_id) \
                .execute()

            # Удаляем URL из data черновика
            draft = SupabaseDB.get_draft(draft_id)
            if draft:
                current_data = draft.get('data', {})
                if 'schemaUrl' in current_data:
                    del current_data['schemaUrl']
                    supabase.table('drafts') \
                        .update({'data': current_data}) \
                        .eq('id', draft_id) \
                        .execute()

            print(f"📐 Удален файл схемы для {draft_id}")
            return True

        except Exception as e:
            print(f"❌ Ошибка удаления файла схемы: {e}")
            return False


# ========== API ЭНДПОИНТЫ ==========

# ========== БАЛАНСИРОВКА ВЕКТОРНЫМ МЕТОДОМ С НАКОПЛЕНИЕМ КВ ==========
@app.route('/api/balancing/vector', methods=['POST'])
def calculate_vector_balancing():
    """
    API для расчета балансировки векторным методом (с фазоизмерительной аппаратурой)
    Возвращает корректирующий груз и угол установки с учетом направления вращения
    """
    try:
        import math

        data = request.get_json()

        if not data:
            return jsonify({
                'success': False,
                'error': 'Нет данных для расчета'
            }), 400

        # Получаем входные параметры
        V0 = data.get('V0')
        phi0 = data.get('phi0')
        Vt = data.get('Vt')
        phi_t = data.get('phi_t')
        P = data.get('P')
        rotation_direction = data.get('rotation_direction', 'cw')  # 'cw' - на себя, 'ccw' - от себя

        # Валидация
        required_fields = ['V0', 'phi0', 'Vt', 'phi_t', 'P']
        for field in required_fields:
            if data.get(field) is None:
                return jsonify({
                    'success': False,
                    'error': f'Поле {field} обязательно для заполнения'
                }), 400

        try:
            V0 = float(V0)
            phi0 = float(phi0)
            Vt = float(Vt)
            phi_t = float(phi_t)
            P = float(P)
        except ValueError:
            return jsonify({
                'success': False,
                'error': 'Все значения должны быть числами'
            }), 400

        if V0 <= 0 or Vt <= 0 or P <= 0:
            return jsonify({
                'success': False,
                'error': 'Амплитуды и масса должны быть положительными числами'
            }), 400

        def to_radians(degrees):
            return degrees * math.pi / 180.0

        def to_degrees(radians):
            return radians * 180.0 / math.pi

        print(f"\n=== Расчет векторным методом ===")
        print(f"V0={V0} мкм, φ0={phi0}°")
        print(f"Vt={Vt} мкм, φt={phi_t}°")
        print(f"P={P} г")
        print(f"Направление вращения: {'на себя (CW)' if rotation_direction == 'cw' else 'от себя (CCW)'}")

        # КОРРЕКТИРОВКА ФАЗЫ В ЗАВИСИМОСТИ ОТ НАПРАВЛЕНИЯ ВРАЩЕНИЯ
        original_phi0 = phi0
        original_phi_t = phi_t

        if rotation_direction == 'ccw':
            # При вращении "от себя" (CCW) фаза отсчитывается в противоположную сторону
            phi0 = (-phi0) % 360
            phi_t = (-phi_t) % 360
            print(f"Скорректированные фазы для CCW: φ0={phi0}°, φt={phi_t}°")

        # Шаг 1: Представляем векторы в комплексной форме
        V0_complex = V0 * (math.cos(to_radians(phi0)) + 1j * math.sin(to_radians(phi0)))
        Vt_complex = Vt * (math.cos(to_radians(phi_t)) + 1j * math.sin(to_radians(phi_t)))

        print(f"V0 комплексный: {V0_complex.real:.3f} + {V0_complex.imag:.3f}i")
        print(f"Vt комплексный: {Vt_complex.real:.3f} + {Vt_complex.imag:.3f}i")

        # Шаг 2: Находим вектор влияния пробного груза W
        W_complex = Vt_complex - V0_complex
        W_magnitude = abs(W_complex)
        W_angle = to_degrees(math.atan2(W_complex.imag, W_complex.real)) % 360

        print(f"Вектор влияния W: {W_magnitude:.3f} @ {W_angle:.1f}°")

        # Шаг 3: Вычисляем удельное влияние пробного груза
        K_complex = W_complex / P
        K_magnitude = abs(K_complex)
        K_angle = to_degrees(math.atan2(K_complex.imag, K_complex.real)) % 360

        print(f"Удельное влияние K: {K_magnitude:.3f} @ {K_angle:.1f}°")

        # Шаг 4: Расчет корректирующего груза
        G_comp_complex = -V0_complex / K_complex
        correction_mass = abs(G_comp_complex)
        correction_angle = to_degrees(math.atan2(G_comp_complex.imag, G_comp_complex.real))

        # Приводим угол к диапазону 0-360
        correction_angle = correction_angle % 360

        # Если вращение "от себя", нужно скорректировать угол обратно
        if rotation_direction == 'ccw':
            correction_angle = (-correction_angle) % 360
            print(f"Скорректированный угол для CCW: {correction_angle}°")

        print(f"Корректирующий груз: {correction_mass:.3f} г @ {correction_angle:.1f}°")

        return jsonify({
            'success': True,
            'result': {
                'correction_mass': round(correction_mass, 3),
                'correction_angle': round(correction_angle),
                'rotation_direction': rotation_direction,
                'debug_info': {
                    'W_magnitude': round(W_magnitude, 3),
                    'W_angle': round(W_angle, 1),
                    'K_magnitude': round(K_magnitude, 3),
                    'K_angle': round(K_angle, 1)
                }
            }
        })

    except Exception as e:
        print(f"❌ Ошибка векторного расчета балансировки: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# ========== НОВЫЕ ЭНДПОИНТЫ ДЛЯ РАБОТЫ СО СХЕМОЙ ==========

@app.route('/api/drafts/<draft_id>/schema', methods=['GET', 'POST', 'DELETE'])
def manage_schema_file(draft_id):
    """Управление файлом электрической схемы"""
    try:
        if request.method == 'GET':
            # Проверяем наличие файла
            file_data = SupabaseDB.get_schema_file(draft_id)

            if file_data:
                file_bytes, filename, content_type = file_data
                return send_file(
                    io.BytesIO(file_bytes),
                    mimetype=content_type,
                    as_attachment=False,
                    download_name=filename
                )
            else:
                return jsonify({
                    'success': False,
                    'error': 'Файл схемы не найден'
                }), 404

        elif request.method == 'POST':
            # Сохраняем файл схемы
            if 'schemaFile' not in request.files:
                return jsonify({
                    'success': False,
                    'error': 'Файл не передан'
                }), 400

            file = request.files['schemaFile']
            if file.filename == '':
                return jsonify({
                    'success': False,
                    'error': 'Файл не выбран'
                }), 400

            success, result = SupabaseDB.save_schema_file(draft_id, file)

            if success:
                return jsonify({
                    'success': True,
                    'message': 'Файл схемы сохранен',
                    'filename': result
                })
            else:
                return jsonify({
                    'success': False,
                    'error': result
                }), 400

        elif request.method == 'DELETE':
            # Удаляем файл схемы
            success = SupabaseDB.delete_schema_file(draft_id)

            if success:
                return jsonify({
                    'success': True,
                    'message': 'Файл схемы удален'
                })
            else:
                return jsonify({
                    'success': False,
                    'error': 'Ошибка удаления файла'
                }), 500

    except Exception as e:
        print(f"❌ Ошибка управления файлом схемы: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/drafts/<draft_id>/schema-info', methods=['GET'])
def get_schema_info(draft_id):
    """Получает информацию о файле схемы (существует ли)"""
    try:
        if supabase is None:
            return jsonify({
                'success': False,
                'error': 'Supabase не инициализирован'
            }), 500

        response = supabase.table('schema_files') \
            .select('filename, uploaded_at') \
            .eq('draft_id', draft_id) \
            .execute()

        if response.data:
            return jsonify({
                'success': True,
                'has_schema': True,
                'filename': response.data[0]['filename'],
                'uploaded_at': response.data[0]['uploaded_at']
            })
        else:
            return jsonify({
                'success': True,
                'has_schema': False
            })

    except Exception as e:
        print(f"❌ Ошибка получения информации о схеме: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/drafts/<draft_id>/request-text', methods=['GET'])
def get_request_text(draft_id):
    """Получает извлеченный текст заявки"""
    try:
        if supabase is None:
            return jsonify({
                'success': False,
                'error': 'Supabase не инициализирован'
            }), 500

        # Получаем текст из базы
        response = supabase.table('request_files') \
            .select('extracted_text, has_text, filename') \
            .eq('draft_id', draft_id) \
            .execute()

        if response.data and response.data[0].get('has_text'):
            return jsonify({
                'success': True,
                'text': response.data[0]['extracted_text'],
                'filename': response.data[0]['filename']
            })
        else:
            return jsonify({
                'success': False,
                'error': 'Текст не найден или не извлечен'
            }), 404

    except Exception as e:
        print(f"❌ Ошибка получения текста заявки: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/drafts/<draft_id>/upload-request', methods=['POST'])
def upload_request_file(draft_id):
    """Загружает файл заявки для черновика"""
    try:
        if 'requestFile' not in request.files:
            return jsonify({
                'success': False,
                'error': 'Файл не передан'
            }), 400

        file = request.files['requestFile']
        if file.filename == '':
            return jsonify({
                'success': False,
                'error': 'Файл не выбран'
            }), 400

        success, result = SupabaseDB.save_request_file(draft_id, file)

        if success:
            return jsonify({
                'success': True,
                'message': 'Файл заявки сохранен',
                'filename': result
            })
        else:
            return jsonify({
                'success': False,
                'error': result
            }), 400

    except Exception as e:
        print(f"❌ Ошибка загрузки файла заявки: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/health', methods=['GET'])
def health_check():
    """Проверка состояния сервера"""
    return jsonify({
        'success': True,
        'status': 'ok',
        'service': 'machine-protocol-generator',
        'supabase_configured': supabase is not None,
        'template_exists': os.path.exists(Config.TEMPLATE_PATH),
        'image_compression': {
            'enabled': True,
            'max_size': Config.IMAGE_MAX_SIZE,
            'quality': Config.IMAGE_QUALITY,
            'format': Config.IMAGE_FORMAT
        },
        'timestamp': datetime.now().isoformat()
    })


@app.route('/api/drafts', methods=['GET'])
def get_drafts():
    """Получает список черновиков"""
    try:
        status_filter = request.args.get('status')
        drafts = SupabaseDB.get_all_drafts(status_filter)

        return jsonify({
            'success': True,
            'drafts': drafts,
            'total': len(drafts),
            'filter': status_filter
        })
    except Exception as e:
        print(f"❌ Ошибка в /api/drafts: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/drafts/<draft_id>', methods=['GET'])
def get_draft(draft_id):
    """Получает конкретный черновик"""
    try:
        draft = SupabaseDB.get_draft(draft_id)

        if draft:
            return jsonify({
                'success': True,
                'draft': draft
            })
        else:
            return jsonify({
                'success': False,
                'error': 'Черновик не найден'
            }), 404
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/save-draft', methods=['POST'])
def save_draft():
    """Сохраняет черновик со сжатыми изображениями"""
    try:
        data = request.form
        images = request.files.getlist('images')

        success, result, error = SupabaseDB.save_draft(data, images)

        if success:
            # Сохраняем файл схемы, если есть
            if 'schemaFile' in request.files:
                schema_file = request.files['schemaFile']
                if schema_file and schema_file.filename:
                    SupabaseDB.save_schema_file(result['id'], schema_file)

            return jsonify({
                'success': True,
                'message': 'Черновик успешно сохранен',
                'draft_id': result['id'],
                'draft_data': result,
                'image_compression': {
                    'max_size': Config.IMAGE_MAX_SIZE,
                    'quality': Config.IMAGE_QUALITY
                }
            })
        else:
            return jsonify({
                'success': False,
                'error': result
            }), 400
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Ошибка сохранения: {str(e)}'
        }), 500


@app.route('/api/drafts/<draft_id>', methods=['PUT'])
def update_draft(draft_id):
    """Обновляет черновик"""
    try:
        data = request.form
        images = request.files.getlist('images')
        success, result = SupabaseDB.update_draft(draft_id, data, images)

        if success:
            # Сохраняем файл схемы, если есть
            if 'schemaFile' in request.files:
                schema_file = request.files['schemaFile']
                if schema_file and schema_file.filename:
                    SupabaseDB.save_schema_file(draft_id, schema_file)

            return jsonify({
                'success': True,
                'message': 'Черновик обновлен',
                'draft_data': result
            })
        else:
            return jsonify({
                'success': False,
                'error': result
            }), 400
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/drafts/<draft_id>/customer', methods=['GET', 'POST', 'PUT'])
def manage_customer_data(draft_id):
    """Управление данными заказчика"""
    try:
        if request.method == 'GET':
            customer_data = SupabaseDB.get_customer_data(draft_id)

            if customer_data:
                return jsonify({
                    'success': True,
                    'customer_data': customer_data
                })
            else:
                return jsonify({
                    'success': False,
                    'error': 'Данные заказчика не найдены'
                }), 404

        elif request.method in ['POST', 'PUT']:
            data = request.get_json()

            if not data:
                return jsonify({
                    'success': False,
                    'error': 'Нет данных для сохранения'
                }), 400

            success, result = SupabaseDB.update_customer_data(draft_id, data)

            if success:
                return jsonify({
                    'success': True,
                    'message': 'Данные заказчика сохранены',
                    'draft_data': result
                })
            else:
                return jsonify({
                    'success': False,
                    'error': result
                }), 400

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/drafts/<draft_id>/images/<filename>')
def get_draft_image(draft_id, filename):
    """Возвращает изображение"""
    try:
        image_data = SupabaseDB.get_image(draft_id, filename)

        if image_data:
            img_bytes, content_type = image_data
            return send_file(
                io.BytesIO(img_bytes),
                mimetype=content_type,
                as_attachment=False,
                download_name=filename
            )
        else:
            return jsonify({
                'success': False,
                'error': 'Изображение не найдено'
            }), 404

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/drafts/<draft_id>/request-file', methods=['GET', 'POST', 'DELETE'])
def manage_request_file(draft_id):
    """Управление файлом заявки"""
    try:
        if request.method == 'GET':
            # Получаем файл заявки
            file_data = SupabaseDB.get_request_file(draft_id)

            if file_data:
                file_bytes, filename, content_type = file_data
                return send_file(
                    io.BytesIO(file_bytes),
                    mimetype=content_type,
                    as_attachment=False,
                    download_name=filename
                )
            else:
                return jsonify({
                    'success': False,
                    'error': 'Файл заявки не найден'
                }), 404

        elif request.method == 'POST':
            # Сохраняем файл заявки
            if 'requestFile' not in request.files:
                return jsonify({
                    'success': False,
                    'error': 'Файл не передан'
                }), 400

            file = request.files['requestFile']
            if file.filename == '':
                return jsonify({
                    'success': False,
                    'error': 'Файл не выбран'
                }), 400

            success, result = SupabaseDB.save_request_file(draft_id, file)

            if success:
                return jsonify({
                    'success': True,
                    'message': 'Файл заявки сохранен',
                    'filename': result
                })
            else:
                return jsonify({
                    'success': False,
                    'error': result
                }), 400

        elif request.method == 'DELETE':
            # Удаляем файл заявки
            success = SupabaseDB.delete_request_file(draft_id)

            if success:
                return jsonify({
                    'success': True,
                    'message': 'Файл заявки удален'
                })
            else:
                return jsonify({
                    'success': False,
                    'error': 'Ошибка удаления файла'
                }), 500

    except Exception as e:
        print(f"❌ Ошибка управления файлом заявки: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/generate-protocol', methods=['POST'])
def generate_protocol():
    """
    Генерирует полный пакет: протокол Excel + все изображения
    и сразу возвращает ZIP архив для скачивания
    """
    try:
        if not os.path.exists(Config.TEMPLATE_PATH):
            return jsonify({
                'success': False,
                'error': f'Шаблон {Config.TEMPLATE_PATH} не найден'
            }), 404

        data = request.form
        draft_id = data.get('draft_id')

        # ПОЛУЧАЕМ ДАННЫЕ ДЛЯ ИМЕН ФАЙЛОВ ЗДЕСЬ, В НАЧАЛЕ
        machine_type = data.get('machineType', '').strip()
        lifting_capacity = data.get('liftingCapacity', '').strip()
        serial_number = data.get('serialNumber', 'unknown').strip().replace(' ', '_')

        # Создаем временную директорию
        with tempfile.TemporaryDirectory() as temp_dir:
            folder_name = generate_folder_name(data)
            protocol_filename = generate_protocol_filename(data)

            # 1. Генерируем протокол Excel из шаблона
            wb = load_workbook(Config.TEMPLATE_PATH)
            ws = wb.active

            # Маппинг полей на ячейки Excel (для первого листа)
            mapping = {
                'workType': 'I1',
                'machineType': 'C3',
                'liftingCapacity': 'D3',
                'serialNumber': 'J3',
                'driveType': 'C5',
                'driveNumber': 'F5',
                'brakeResistor': 'E7',
                'resistorCount': 'H7',
                'electricMotor': 'D9',
                'motorNumber': 'G9',
                'EnginePower': 'K9',
                'angleSensor': 'D11',
                'angleSensorNumber': 'G11',
                'speedSensorNumber': 'K11',
                'leftVibrationSensor': 'D15',
                'leftSensitivity': 'G15',
                'leftSensorNumber': 'I15',
                'rightVibrationSensor': 'D16',
                'rightSensitivity': 'G16',
                'rightSensorNumber': 'I16',
                'measuringDevice': 'E18',
                'measuringDeviceNumber': 'G18',
                'signalProcessor': 'E20',
                'signalProcessorNumber': 'G20'
            }

            for field, cell in mapping.items():
                if field in data and data[field]:
                    ws[cell] = data[field]

            # Записываем примечания на второй лист в объединенную ячейку B2:G52
            if 'notes' in data and data['notes']:
                ws2 = wb.worksheets[1]  # Второй лист (индекс 1)
                # Получаем мастер-ячейку объединенного диапазона
                merged_range = 'B2:G52'
                if merged_range in ws2.merged_cells:
                    # Находим мастер-ячейку (первую в объединенном диапазоне)
                    for merged_cell in ws2.merged_cells:
                        if merged_cell.coord == merged_range:
                            ws2.cell(row=merged_cell.min_row, column=merged_cell.min_col).value = data['notes']
                            break
                else:
                    # Если диапазон не объединен, записываем в B2
                    ws2['B2'] = data['notes']

            # Сохраняем протокол
            protocol_path = os.path.join(temp_dir, protocol_filename)
            wb.save(protocol_path)
            wb.close()  # Явно закрываем workbook

            # 2. Создаем папку для изображений
            images_dir = os.path.join(temp_dir, f"{folder_name}_images")
            os.makedirs(images_dir, exist_ok=True)

            # 3. Сохраняем изображения из формы
            images = request.files.getlist('images')
            saved_images = []

            for img in images:
                if img and allowed_file(img.filename):
                    filename = secure_filename(img.filename)
                    img.seek(0)
                    img_data = img.read()

                    # Сжимаем изображение
                    compressed_data = compress_image(
                        img_data,
                        max_size=Config.IMAGE_MAX_SIZE,
                        quality=Config.IMAGE_QUALITY
                    )

                    # Сохраняем во временную папку
                    img_path = os.path.join(images_dir, filename)
                    with open(img_path, 'wb') as f:
                        f.write(compressed_data)
                    saved_images.append(filename)
                    print(f"📸 Сохранено изображение из формы: {filename}")

            # 4. Если есть draft_id, загружаем дополнительные изображения из БД
            if draft_id:
                draft = SupabaseDB.get_draft(draft_id)
                if draft and draft.get('image_files'):
                    for img_filename in draft['image_files']:
                        if img_filename not in saved_images:
                            image_data = SupabaseDB.get_image(draft_id, img_filename)
                            if image_data:
                                img_bytes, _ = image_data
                                img_path = os.path.join(images_dir, img_filename)
                                with open(img_path, 'wb') as f:
                                    f.write(img_bytes)
                                saved_images.append(img_filename)
                                print(f"📸 Сохранено изображение из БД: {img_filename}")

            # 5. Создаем информационный файл README.txt
            info_content = f"""
ПРОТОКОЛ

=====================================


1. ОСНОВНАЯ ИНФОРМАЦИЯ
-----------------------
Тип станка:            {data.get('machineType', 'Не указан')}
Грузоподъемность:      {data.get('liftingCapacity', 'Не указана')}
Заводской номер:       {data.get('serialNumber', 'Не указан')}
Вид работ:            {data.get('workType', 'Не указан')}
Заказчик:             {data.get('customer', 'Не указан')}
Статус:               {data.get('machineStatus', 'Сборка')}
Плановая отгрузка:    {data.get('shippingDate', 'Не указана')}

2. ПРИВОДНАЯ СИСТЕМА
--------------------
Тип привода:          {data.get('driveType', 'Не указан')}
Номер привода:        {data.get('driveNumber', 'Не указан')}
Тормозной резистор:   {data.get('brakeResistor', 'Не указан')}
Кол-во резисторов:    {data.get('resistorCount', '-')}

3. ЭЛЕКТРОДВИГАТЕЛЬ
-------------------
Тип двигателя:        {data.get('electricMotor', 'Не указан')}
Мощность:             {data.get('EnginePower', 'Не указана')}
Номер двигателя:      {data.get('motorNumber', 'Не указан')}

4. ДАТЧИКИ
-----------
Датчик угла:          {data.get('angleSensor', 'Не указан')}
Номер датчика угла:   {data.get('angleSensorNumber', 'Не указан')}
Номер отметчика:      {data.get('speedSensorNumber', 'Не указан')}

5. ДАТЧИКИ ВИБРАЦИИ
-------------------
Левый датчик:         {data.get('leftVibrationSensor', 'Не указан')}
Чувствительность:     {data.get('leftSensitivity', '-')}
Номер датчика:        {data.get('leftSensorNumber', 'Не указан')}

Правый датчик:        {data.get('rightVibrationSensor', 'Не указан')}
Чувствительность:     {data.get('rightSensitivity', '-')}
Номер датчика:        {data.get('rightSensorNumber', 'Не указан')}

6. ИЗМЕРИТЕЛЬНОЕ ОБОРУДОВАНИЕ
-----------------------------
Измерительный прибор: {data.get('measuringDevice', 'Не указан')}
Номер прибора:        {data.get('measuringDeviceNumber', 'Не указан')}
Блок обработки:       {data.get('signalProcessor', 'Не указан')}
Номер блока:          {data.get('signalProcessorNumber', 'Не указан')}


8. ПРИМЕЧАНИЯ
-------------
{data.get('notes', 'Нет')}

=====================================
Конец протокола
            """

            info_path = os.path.join(temp_dir, "README.txt")
            with open(info_path, 'w', encoding='utf-8') as f:
                f.write(info_content.strip())

            # 6. Создаем ZIP архив
            zip_filename = generate_zip_filename(data)
            zip_path = os.path.join(temp_dir, zip_filename)

            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                # Добавляем протокол Excel
                zf.write(protocol_path, protocol_filename)

                # Добавляем README с информацией
                zf.write(info_path, "README.txt")

                # Добавляем все изображения в папку images
                for img_file in os.listdir(images_dir):
                    img_full_path = os.path.join(images_dir, img_file)
                    zf.write(img_full_path, f"images/{img_file}")

            print(f"✅ ZIP архив создан: {zip_filename}, размер: {os.path.getsize(zip_path)} байт")

            # 7. Читаем ZIP файл в память, чтобы закрыть его до удаления временной директории
            with open(zip_path, 'rb') as f:
                zip_data = f.read()

            # 8. Отправляем ZIP файл клиенту
            return send_file(
                io.BytesIO(zip_data),
                mimetype='application/zip',
                as_attachment=True,
                download_name=zip_filename
            )

    except Exception as e:
        print(f"❌ Ошибка генерации пакета: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'Ошибка генерации протокола: {str(e)}'
        }), 500


@app.route('/api/history', methods=['GET'])
def get_history():
    """
    Получает историю изменений всех черновиков
    Поддерживает пагинацию и фильтрацию по дате
    """
    try:
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 50))
        date = request.args.get('date')

        if date:
            history = SupabaseDB.get_history_by_date(date, page, per_page)
        else:
            history = SupabaseDB.get_history(page=page, per_page=per_page)

        return jsonify({
            'success': True,
            'history': history['items'],
            'pagination': {
                'page': history['page'],
                'per_page': history['per_page'],
                'total': history['total'],
                'total_pages': history['total_pages']
            }
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/history/<draft_id>', methods=['GET'])
def get_draft_history(draft_id):
    """
    Получает историю изменений конкретного черновика
    """
    try:
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 50))

        history = SupabaseDB.get_history(draft_id, page, per_page)

        return jsonify({
            'success': True,
            'draft_id': draft_id,
            'history': history['items'],
            'pagination': {
                'page': history['page'],
                'per_page': history['per_page'],
                'total': history['total'],
                'total_pages': history['total_pages']
            }
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# ========== ЭНДПОИНТЫ ДЛЯ УПРАВЛЕНИЯ ПРОШИВКАМИ ==========

@app.route('/api/updates', methods=['GET'])
def get_updates():
    """Получает все прошивки и утилиты"""
    try:
        if supabase is None:
            return jsonify({
                'success': False,
                'error': 'Supabase не инициализирован'
            }), 500

        section = request.args.get('section')

        query = supabase.table('updates').select('*')
        if section:
            query = query.eq('section', section)

        response = query.order('sort_order', desc=False).order('created_at', desc=True).execute()

        # Получаем дату последнего обновления
        meta_response = supabase.table('updates_meta') \
            .select('value') \
            .eq('key', 'last_update') \
            .execute()

        last_update = meta_response.data[0]['value'] if meta_response.data else '24.02.2026'

        # Группируем по секциям для удобства
        grouped = {
            'sapphire': [],
            'yashma': [],
            'stsh': [],
            'external_utils': [],
            'internal_utils': [],
            'docs': []
        }

        for item in response.data:
            section = item.get('section')
            if section in grouped:
                grouped[section].append({
                    'id': item.get('id'),
                    'title': item.get('title'),
                    'description': item.get('description'),
                    'date': item.get('date'),
                    'file': item.get('file_path'),
                    'badge': item.get('badge'),
                    'sort_order': item.get('sort_order', 0)
                })

        return jsonify({
            'success': True,
            'updates': grouped,
            'last_update': last_update
        })

    except Exception as e:
        print(f"❌ Ошибка получения прошивок: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/updates', methods=['POST'])
def create_update():
    """Создает новую запись о прошивке"""
    try:
        if supabase is None:
            return jsonify({
                'success': False,
                'error': 'Supabase не инициализирован'
            }), 500

        data = request.get_json()

        if not data:
            return jsonify({
                'success': False,
                'error': 'Нет данных для сохранения'
            }), 400

        # Валидация
        required_fields = ['section', 'title', 'file']
        for field in required_fields:
            if field not in data or not data[field]:
                return jsonify({
                    'success': False,
                    'error': f'Поле {field} обязательно'
                }), 400

        # Подготавливаем данные
        update_data = {
            'section': data.get('section'),
            'title': data.get('title'),
            'description': data.get('description', ''),
            'date': data.get('date', ''),
            'file_path': data.get('file'),
            'badge': data.get('badge', 'ZIP'),
            'sort_order': data.get('sort_order', 0),
            'updated_at': datetime.now().isoformat()
        }

        # Вставляем в БД
        response = supabase.table('updates').insert(update_data).execute()

        # Обновляем дату последнего изменения
        supabase.table('updates_meta') \
            .update({'value': datetime.now().strftime('%d.%m.%Y'), 'updated_at': datetime.now().isoformat()}) \
            .eq('key', 'last_update') \
            .execute()

        return jsonify({
            'success': True,
            'message': 'Запись успешно создана',
            'update': response.data[0] if response.data else None
        })

    except Exception as e:
        print(f"❌ Ошибка создания записи: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/updates/<int:update_id>', methods=['PUT'])
def update_update(update_id):
    """Обновляет существующую запись"""
    try:
        if supabase is None:
            return jsonify({
                'success': False,
                'error': 'Supabase не инициализирован'
            }), 500

        data = request.get_json()

        if not data:
            return jsonify({
                'success': False,
                'error': 'Нет данных для обновления'
            }), 400

        # Подготавливаем данные для обновления
        update_data = {
            'title': data.get('title'),
            'description': data.get('description', ''),
            'date': data.get('date', ''),
            'file_path': data.get('file'),
            'badge': data.get('badge', 'ZIP'),
            'sort_order': data.get('sort_order', 0),
            'updated_at': datetime.now().isoformat()
        }

        # Убираем None значения
        update_data = {k: v for k, v in update_data.items() if v is not None}

        # Обновляем в БД
        response = supabase.table('updates') \
            .update(update_data) \
            .eq('id', update_id) \
            .execute()

        # Обновляем дату последнего изменения
        supabase.table('updates_meta') \
            .update({'value': datetime.now().strftime('%d.%m.%Y'), 'updated_at': datetime.now().isoformat()}) \
            .eq('key', 'last_update') \
            .execute()

        return jsonify({
            'success': True,
            'message': 'Запись успешно обновлена',
            'update': response.data[0] if response.data else None
        })

    except Exception as e:
        print(f"❌ Ошибка обновления записи: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/drafts/<draft_id>/delete', methods=['POST', 'DELETE'])
def delete_draft(draft_id):
    """
    Удаляет черновик со всеми связанными данными
    Доступен только через специальный URL параметр для безопасности
    """
    try:
        # Проверяем специальный секретный ключ для защиты от случайного удаления
        confirm = request.args.get('confirm')

        if confirm != 'yes':
            return jsonify({
                'success': False,
                'error': 'Для удаления требуется подтверждение'
            }), 400

        # Выполняем удаление
        success, message = SupabaseDB.delete_draft(draft_id)

        if success:
            return jsonify({
                'success': True,
                'message': message,
                'draft_id': draft_id
            })
        else:
            return jsonify({
                'success': False,
                'error': message
            }), 404 if message == "Черновик не найден" else 500

    except Exception as e:
        print(f"❌ Ошибка в эндпоинте удаления: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/updates/<int:update_id>', methods=['DELETE'])
def delete_update(update_id):
    """Удаляет запись"""
    try:
        if supabase is None:
            return jsonify({
                'success': False,
                'error': 'Supabase не инициализирован'
            }), 500

        # Удаляем из БД
        response = supabase.table('updates') \
            .delete() \
            .eq('id', update_id) \
            .execute()

        # Обновляем дату последнего изменения
        supabase.table('updates_meta') \
            .update({'value': datetime.now().strftime('%d.%m.%Y'), 'updated_at': datetime.now().isoformat()}) \
            .eq('key', 'last_update') \
            .execute()

        return jsonify({
            'success': True,
            'message': 'Запись успешно удалена'
        })

    except Exception as e:
        print(f"❌ Ошибка удаления записи: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/updates/reorder', methods=['POST'])
def reorder_updates():
    """Изменяет порядок сортировки записей"""
    try:
        if supabase is None:
            return jsonify({
                'success': False,
                'error': 'Supabase не инициализирован'
            }), 500

        data = request.get_json()
        orders = data.get('orders', [])

        if not orders:
            return jsonify({
                'success': False,
                'error': 'Нет данных для сортировки'
            }), 400

        # Обновляем порядок для каждой записи
        for item in orders:
            supabase.table('updates') \
                .update({'sort_order': item['order']}) \
                .eq('id', item['id']) \
                .execute()

        return jsonify({
            'success': True,
            'message': 'Порядок сортировки обновлен'
        })

    except Exception as e:
        print(f"❌ Ошибка обновления сортировки: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# ========== СТАТИЧЕСКИЕ ФАЙЛЫ ==========

@app.route('/')
def serve_main():
    """Главная страница"""
    return send_from_directory(app.static_folder, 'main.html')


@app.route('/add-draft')
def serve_add_draft():
    """Страница добавления станка"""
    return send_from_directory(app.static_folder, 'add_draft.html')


@app.route('/view-machine.html')
def serve_view_machine():
    """Страница просмотра станка"""
    return send_from_directory(app.static_folder, 'view_machine.html')


@app.route('/static/<path:path>')
def serve_static(path):
    """Статические файлы"""
    return send_from_directory('static', path)


@app.route('/favicon.ico')
def favicon():
    """Иконка сайта"""
    return send_from_directory('static', 'favicon.ico'), 404


# ========== ОБРАБОТЧИКИ ОШИБОК ==========

@app.errorhandler(404)
def not_found(error):
    return jsonify({
        'success': False,
        'error': 'Эндпоинт не найден'
    }), 404


@app.errorhandler(500)
def internal_error(error):
    return jsonify({
        'success': False,
        'error': 'Внутренняя ошибка сервера'
    }), 500


# ========== ЗАПУСК ПРИЛОЖЕНИЯ ==========

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 10000))
    app.run(host='0.0.0.0', port=port, debug=False)