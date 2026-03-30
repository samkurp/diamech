"""
Утилиты для работы с документами (DOCX, DOC)
"""
import io
import re
from docx import Document
import olefile


def extract_text_from_docx(file_data):
    """Извлекает текст из DOCX файла"""
    try:
        doc = Document(io.BytesIO(file_data))
        text_parts = []
        
        # Извлекаем текст из параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        # Извлекаем текст из заголовков
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(f"[Заголовок] {paragraph.text.strip()}")
            
            footer = section.footer
            for paragraph in footer.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(f"[Подвал] {paragraph.text.strip()}")
        
        return '\n\n'.join(text_parts) if text_parts else None
        
    except Exception as e:
        print(f"❌ Ошибка извлечения текста из DOCX: {e}")
        return None


def extract_text_from_doc(file_data):
    """Извлекает текст из DOC файла (старый формат)"""
    try:
        ole = olefile.OleFileIO(io.BytesIO(file_data))
        
        if not ole.exists('WordDocument'):
            return None
        
        # Пытаемся извлечь текст из разных потоков
        text_parts = []
        
        # Основной текст
        if ole.exists('WordDocument'):
            stream = ole.openstream('WordDocument')
            data = stream.read()
            
            # Пытаемся декодировать текст
            try:
                # Пробуем разные кодировки
                for encoding in ['utf-16', 'utf-8', 'latin-1', 'cp1251']:
                    try:
                        text = data.decode(encoding, errors='ignore')
                        # Очищаем от нечитаемых символов
                        text = re.sub(r'[^\x20-\x7Eа-яА-ЯёЁ\s\n\r]', ' ', text)
                        # Разбиваем на строки
                        lines = [line.strip() for line in text.split('\n') if line.strip()]
                        text_parts.extend(lines)
                        break
                    except:
                        continue
            except:
                pass
        
        # Пытаемся получить текст из потока 1Table
        if ole.exists('1Table'):
            stream = ole.openstream('1Table')
            data = stream.read()
            try:
                text = data.decode('utf-16', errors='ignore')
                text = re.sub(r'[^\x20-\x7Eа-яА-ЯёЁ\s\n\r]', ' ', text)
                lines = [line.strip() for line in text.split('\n') if line.strip()]
                text_parts.extend(lines)
            except:
                pass
        
        ole.close()
        
        return '\n\n'.join(text_parts) if text_parts else None
        
    except Exception as e:
        print(f"❌ Ошибка извлечения текста из DOC: {e}")
        return None


def extract_text_from_file(file_data, filename):
    """Определяет тип файла и извлекает текст"""
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    
    if ext == 'docx':
        return extract_text_from_docx(file_data)
    elif ext == 'doc':
        return extract_text_from_doc(file_data)
    else:
        return None
